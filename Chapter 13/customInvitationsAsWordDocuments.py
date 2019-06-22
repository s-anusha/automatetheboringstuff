#! python3
# customInvitationsAsWordDocuments.py
'''
Say you have a text file of guest names. This guests.txt file has one name per line, as follows:
Prof. Plum
Miss Scarlet
Col. Mustard
Al Sweigart
Robocop
Write a program that would generate a Word document with custom invitations that look like Figure 13-11.
Since Python-Docx can use only those styles that already exist in the Word document, you will have to first add these styles to a blank Word file and then open that file with Python-Docx. There should be one invitation per page in the resulting Word document, so call add_break() to add a page break after the last paragraph of each invitation. This way, you will need to open only one Word document to print all of the invitations at once.
This program uses preexisting styles.
'''
# Usage: python customInvitationsAsWordDocuments.py guest_list

import sys
import docx

if len(sys.argv) < 2:
    print('Usage: python customInvitationAsWordDocuments.py guests.txt')
    sys.exit()

doc = docx.Document()
guestList = sys.argv[1]
file = open(guestList, 'r')

for line in file:
    doc.add_heading('It would be a pleasure to have the company of', 0)
    guest =  doc.add_paragraph(line)
    guest.style = 'Subtitle'
    doc.add_heading('at 11010 Memory Lane on the Evening of', 1)
    date = doc.add_paragraph('April 1st')
    date.runs[0].underline = True
    time = doc.add_paragraph('at 7 o\'clock')
    time.runs[0].underline = True
    time.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

file.close()
doc.save('customInvitationsAsWordDocuments.docx')
